import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_resume():
    doc = Document()
    
    # 1. Page Margins (0.5 inches all around for maximum high-density single-column space)
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        
    # Helper to add a clean, horizontal bottom border to paragraphs (for section headings)
    def add_bottom_border(paragraph, color_hex="1F4E79", size=12):
        pPr = paragraph._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(size)) # 12 = 1.5 pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)

    # Helper to add section headers
    def add_section_header(title):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        r = p.add_run(title)
        r.font.name = 'Arial'
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79) # Primary Navy
        r.bold = True
        
        add_bottom_border(p, "1F4E79", 12)
        return p

    # Helper to add experience/project/education headers with tab stops
    def add_header_with_tabs(left_runs, right_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(1.5)
        p.paragraph_format.keep_with_next = True
        
        # Add tab stop at the right margin (7.5 inches for 0.5" margins on an 8.5" wide page)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        
        for text, bold, italic in left_runs:
            r = p.add_run(text)
            r.font.name = 'Arial'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            r.bold = bold
            r.italic = italic
            
        r_tab = p.add_run("\t")
        
        r_right = p.add_run(right_text)
        r_right.font.name = 'Arial'
        r_right.font.size = Pt(10)
        r_right.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r_right.bold = True
        return p

    # Helper for adding job role paragraph
    def add_role_paragraph(role_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.keep_with_next = True
        
        r = p.add_run(role_text)
        r.font.name = 'Arial'
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0x59, 0x59, 0x59) # Cool Grey / Secondary Accent
        r.italic = True
        return p

    # Helper for adding highly customized bullet points (narrow hang-indent)
    def add_bullet_point(bold_prefix, text_content):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.05
        
        # Custom elegant navy bullet symbol
        r_bullet = p.add_run("\u2022  ")
        r_bullet.font.name = 'Arial'
        r_bullet.font.size = Pt(9.5)
        r_bullet.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r_bullet.bold = True
        
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Arial'
            r_bold.font.size = Pt(9.5)
            r_bold.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            r_bold.bold = True
            
        r_text = p.add_run(text_content)
        r_text.font.name = 'Arial'
        r_text.font.size = Pt(9.5)
        r_text.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        return p

    # ==================== HEADER SECTION ====================
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_before = Pt(0)
    p_name.paragraph_format.space_after = Pt(2)
    p_name.paragraph_format.keep_with_next = True
    
    r_name = p_name.add_run("ARIHANT KUMAR JAIN")
    r_name.font.name = 'Arial'
    r_name.font.size = Pt(18)
    r_name.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r_name.bold = True
    
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_before = Pt(0)
    p_contact.paragraph_format.space_after = Pt(8)
    p_contact.paragraph_format.keep_with_next = True
    
    contacts = [
        ("Email: ", "arihantjainn04@gmail.com"),
        ("Phone: ", "+91 8939258516"),
        ("LinkedIn: ", "linkedin.com/in/arihant-jain-ai"),
        ("GitHub: ", "github.com/techwizAJ"),
        ("Portfolio: ", "arihant-tech.info")
    ]
    
    for i, (label, val) in enumerate(contacts):
        r_lbl = p_contact.add_run(label)
        r_lbl.font.name = 'Arial'
        r_lbl.font.size = Pt(9.5)
        r_lbl.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
        r_lbl.bold = True
        
        r_val = p_contact.add_run(val)
        r_val.font.name = 'Arial'
        r_val.font.size = Pt(9.5)
        r_val.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        
        if i == 1:
            p_contact.add_run("\n")
        elif i < len(contacts) - 1:
            r_sep = p_contact.add_run("  |  ")
            r_sep.font.name = 'Arial'
            r_sep.font.size = Pt(9.5)
            r_sep.font.color.rgb = RGBColor(0x59, 0x59, 0x59)

    # ==================== PROFILE SUMMARY ====================
    add_section_header("PROFILE SUMMARY")
    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_before = Pt(0)
    p_summary.paragraph_format.space_after = Pt(6)
    p_summary.paragraph_format.line_spacing = 1.1
    
    summary_text = (
        "High-performance Senior / Staff Software Engineer (L6 equivalent) with 7+ years of hands-on experience "
        "architecting, scaling, and securing distributed ledger, transactional, and next-generation AI platforms. "
        "Recognized expert in optimizing p99 latency in ultra-high throughput environments, resolving complex distributed "
        "state challenges, and integrating autonomous multi-agent workflows as high-scale system capabilities. Proven track "
        "record of spearheading zero-downtime microservice migrations, designing resilient transactional workflows ($10B+ daily volume), "
        "and cutting operational latency by up to 99%. Expert in Java/Spring ecosystem, Kafka event-driven backbones, caching topologies, "
        "and robust system-level consistency models."
    )
    r_sum = p_summary.add_run(summary_text)
    r_sum.font.name = 'Arial'
    r_sum.font.size = Pt(9.5)
    r_sum.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ==================== CORE SKILLS ====================
    add_section_header("CORE SKILLS")
    skills = [
        ("Languages: ", "Java (Core, Concurrency, Multithreading), Python, C++, Node.js, SQL"),
        ("Backend Frameworks: ", "Spring Boot, Spring Cloud, Spring Security, Spring WebFlux (Reactive Stack)"),
        ("Distributed Systems & Storage: ", "Apache Kafka, Redis (Cache-Aside, Distributed Locks), MongoDB, MySQL, Elasticsearch, Azure Service Bus"),
        ("Systems Architecture: ", "Distributed Transactions (Saga Orchestration, Event Sourcing, ACID), Consistency Models (Eventual/Strong Consistency, Linearizability), High Availability, Idempotency Engineering, Fault Isolation (Resilience4j Circuit Breakers, Bulkheads, Rate Limiters), Observability (OpenTelemetry, Prometheus, ELK Stack, Zipkin, JProfiler)"),
        ("AI & Next-Gen Systems: ", "Multi-Agent Systems, LangGraph (Orchestration & State Management), Retrieval-Augmented Generation (RAG) Pipelines, Knowledge Graphs, Semantic Vector Search, Claude / OpenAI LLM Integration"),
        ("Cloud & DevOps: ", "AWS, Azure, Docker, Kubernetes, Terraform, Jenkins, CI/CD, Liquibase Schema Migration")
    ]
    for category, items in skills:
        p_skill = doc.add_paragraph()
        p_skill.paragraph_format.space_before = Pt(0)
        p_skill.paragraph_format.space_after = Pt(2.5)
        p_skill.paragraph_format.line_spacing = 1.05
        
        r_bullet = p_skill.add_run("\u2022  ")
        r_bullet.font.name = 'Arial'
        r_bullet.font.size = Pt(9.5)
        r_bullet.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
        r_bullet.bold = True
        
        r_cat = p_skill.add_run(category)
        r_cat.font.name = 'Arial'
        r_cat.font.size = Pt(9.5)
        r_cat.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r_cat.bold = True
        
        r_items = p_skill.add_run(items)
        r_items.font.name = 'Arial'
        r_items.font.size = Pt(9.5)
        r_items.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ==================== WORK EXPERIENCE ====================
    add_section_header("WORK EXPERIENCE")
    
    # 1. WISSEN TECHNOLOGY
    wissen_left = [
        ("WISSEN TECHNOLOGY", True, False),
        (" | Client: ", False, False),
        ("MORGAN STANLEY", True, False),
        (" | Mumbai, India", False, False)
    ]
    add_header_with_tabs(wissen_left, "June 2025 – Present")
    add_role_paragraph("Senior Software Engineer")
    
    add_bullet_point("Agentic Flow Accelerator Platform: ", 
                     "Architected and deployed an enterprise-grade meta-orchestration platform utilizing autonomous multi-agent networks (via LangGraph and custom semantic state trees) to automate institutional-grade market onboarding (legal entities, depositories, agent accounts). Built event-driven compilation engines that dynamically generate schema migrations (Liquibase) and pull requests, compressing the onboarding pipeline from 3 weeks to ~10 minutes (a 99.9% operational efficiency gain) with zero schema drift.")
    
    add_bullet_point("Agentic Reconciliation Intelligence System: ", 
                     "Engineered a distributed data-aggregation and log-analysis engine to index heterogeneous, multi-system database and legacy mainframe ledger streams (processing billions of events). Integrated a custom RAG pipeline and dynamic Knowledge Graphs into a hierarchical, multi-agent diagnostic framework. Automated break detection and root-cause analysis, compressing cross-border trade desk investigations from weeks to <5 minutes.")
    
    add_bullet_point("Global Market Scoping Engine: ", 
                     "Redesigned Morgan Stanley’s legacy trade-scoping rules system, migrating from a brittle, rules-based structure prone to validation failures to a market-driven policy matrix. Reduced overall rule-base complexity by 85%, eliminating edge cases that previously exposed the firm to millions in transactional settlement risk.")
    
    add_bullet_point("High-Availability Settlement Pipelines: ", 
                     "Led the distributed systems design for the Asia-Pacific trade settlement infrastructure, safely processing 1.5M+ transactions daily ($10B+ transactional volume). Implemented distributed transactional consensus, Saga orchestration patterns, and idempotent ingestion filters to guarantee strict 'exactly-once' processing under extreme market volatility.")

    # 2. TATA CONSULTANCY SERVICES (TCS)
    tcs_left = [
        ("TATA CONSULTANCY SERVICES (TCS)", True, False),
        (" | Client: ", False, False),
        ("ERICSSON", True, False),
        (" | Pune, India", False, False)
    ]
    add_header_with_tabs(tcs_left, "April 2022 – June 2025")
    add_role_paragraph("Systems Engineer")
    
    add_bullet_point("Core Microservice Refactoring & Latency Optimization: ", 
                     "Spearheaded the architectural redesign of a highly degraded core telecommunications provisioning microservice. Replaced synchronous, blocking HTTP polling with a high-throughput, event-driven data ingestion pipeline using Apache Kafka and a distributed Redis cache-aside topology. Slashed p99 latency by 99.6% (from 12 seconds to 40 milliseconds) while enforcing strict eventual consistency models between MongoDB and relational datastores.")
    
    add_bullet_point("High-Throughput Messaging & Scaling: ", 
                     "Designed and deployed a highly scalable, partitioned event-driven messaging backbone using Kafka, configured with custom partitioners, consumer group scaling, and backpressure handling. Achieved throughput scaling of 100k+ RPS during peak telemetry traffic bursts without message loss.")
    
    add_bullet_point("Asynchronous Execution & Concurrency: ", 
                     "Re-engineered legacy synchronous APIs into non-blocking, reactive asynchronous streams (using Spring WebFlux/CompletableFuture), improving concurrent connection handling by 5x under peak telecommunication workloads.")
    
    add_bullet_point("Fault Isolation & Resilience: ", 
                     "Implemented system resilience patterns utilizing Resilience4j (Circuit Breakers, Bulkheads, and Rate Limiters) to prevent cascading failures across downstream provisioning networks. Integrated Azure Active Directory (AD) with OIDC/OAuth 2.0 and Spring Security to enforce enterprise-grade API security.")

    # 3. HIGHRADIUS TECHNOLOGIES
    hr_left = [
        ("HIGHRADIUS TECHNOLOGIES", True, False),
        (" | Hyderabad, India", False, False)
    ]
    add_header_with_tabs(hr_left, "January 2019 – March 2022")
    add_role_paragraph("Technical Consultant (Growth Path: Intern to TC)")
    
    add_bullet_point("Promotion Timeline: ", 
                     "Promoted 4 times within a 3-year tenure, progressing from Software Development Intern to Technical Consultant: Technical Consultant (Oct 2021 – Mar 2022) \u2022 Associate Technical Consultant (Jul 2020 – Nov 2021) \u2022 Associate Software Engineer (May 2019 – Jun 2020) \u2022 Software Development Intern (Jan 2019 – Apr 2019).")
    
    add_bullet_point("Configurable Rule-Engine Framework: ", 
                     "Architected and built a highly configurable, high-performance rule-engine framework to parse and match incoming accounts receivable data. Elevated straight-through payment reconciliation automation from 20% to 80% for Fortune 500 enterprise clients, reducing manual accounts-receivable aging cycles by 70%.")
    
    add_bullet_point("Financial File Ingestion Pipeline: ", 
                     "Engineered a robust, high-volume financial parsing and ingestion pipeline supporting standard global banking formats (BAI2 and MT940). Designed streaming parsers that processed 500MB+ banking statement files containing millions of transactions with strict transactional isolation and duplicate detection (idempotent ledger writes).")
    
    add_bullet_point("Distributed Observability & Performance Tuning: ", 
                     "Designed and established an enterprise-wide observability framework integrating distributed tracing (ELK Stack, Prometheus, Zipkin) and structured logging. Enabled proactive performance bottleneck identification, reducing mean-time-to-detection (MTTD) of anomalies from 4 hours to under 5 minutes across multi-tenant SaaS environments.")
    
    add_bullet_point("Technical Leadership: ", 
                     "Mentored a cross-functional team of 10+ junior software engineers, standardizing code quality metrics, CI/CD automated test suites (JUnit/Mockito), and robust system design best practices.")

    # ==================== KEY PROJECTS ====================
    add_section_header("KEY PROJECTS")
    
    p1_left = [
        ("BoardroomIQ AI", True, False),
        (" | Solo-Built AI Strategy Platform", False, False)
    ]
    add_header_with_tabs(p1_left, "Live Demo: boardroomiq-ai.com")
    add_bullet_point("System Overview: ", 
                     "A high-performance strategy simulation platform featuring autonomous multi-agent boardrooms.")
    add_bullet_point("Distributed & AI Architecture: ", 
                     "Engineered using a hierarchical agent architecture (orchestrated via LangGraph and powered by Claude LLM API) where individual personas (CFO, Bear, Bull, Activist) debate corporate actions using deep Retrieval-Augmented Generation (RAG) context.")
    add_bullet_point("Systems Engineering: ", 
                     "Leveraged Redis for distributed state management and session caching, ensuring low-latency agent turn-taking and context window compression. Designed an asynchronous job queue using Node.js to manage concurrency spikes during heavy simulated debates. Built dynamic evaluation nodes that calculate, track, and score user strategy performance across 8 dimensions in real-time.")

    p2_left = [
        ("Treasury Infrastructure Platform", True, False),
        (" | Open-Source Distributed Financial Infrastructure", False, False)
    ]
    add_header_with_tabs(p2_left, "Source: github.com/techwizAJ/TreasuryApplication")
    add_bullet_point("System Overview: ", 
                     "A high-throughput, enterprise-grade financial ledger and cash-visibility engine.")
    add_bullet_point("Architecture: ", 
                     "Engineered with Spring Boot and MySQL to handle distributed transactional integrity across multi-region accounting ledgers.")
    add_bullet_point("Systems Engineering: ", 
                     "Implemented strict transactional auditing, double-entry bookkeeping constraints, Saga orchestration to handle payment settlements, and idempotent API endpoints to guarantee zero double-spend or duplicate billing events under network failures.")

    # ==================== EDUCATION ====================
    add_section_header("EDUCATION")
    
    def add_education_item(left_runs, right_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.keep_with_next = True
        
        p.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        
        for text, bold, italic in left_runs:
            r = p.add_run(text)
            r.font.name = 'Arial'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            r.bold = bold
            r.italic = italic
            
        r_tab = p.add_run("\t")
        
        r_right = p.add_run(right_text)
        r_right.font.name = 'Arial'
        r_right.font.size = Pt(9.5)
        r_right.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        r_right.bold = True
        return p

    srm_left = [
        ("SRM University, Chennai", True, False),
        (" — ", False, False),
        ("Bachelor of Technology (B.Tech.) in Computer Science and Engineering", False, True),
        (" (2019)", False, False)
    ]
    add_education_item(srm_left, "Performance: 85.09% (Distinction)")
    
    chaitanya_left = [
        ("Sri Chaitanya Junior College, Hyderabad", True, False),
        (" — ", False, False),
        ("Higher Secondary Certificate (12th Grade)", False, True),
        (" (2015)", False, False)
    ]
    add_education_item(chaitanya_left, "Performance: 90.3%")
    
    slate_left = [
        ("Slate The School, Hyderabad", True, False),
        (" — ", False, False),
        ("Secondary School Certificate (10th Grade)", False, True),
        (" (2013)", False, False)
    ]
    add_education_item(slate_left, "Performance: 9.3 GPA")

    # Ensure output directory exists and save document
    output_path = "/home/arihantk/AI/arihant-portfolio/210809CS5486362.docx"
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Resume saved successfully to {output_path}")

if __name__ == "__main__":
    create_resume()
